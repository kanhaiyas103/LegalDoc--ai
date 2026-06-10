from io import BytesIO
from typing import Any

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def _flatten(result: dict[str, Any]) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    for key, value in result.items():
        label = key.replace("_", " ").title()
        rows.append((label, str(value)))
    return rows


def analysis_pdf(*, title: str, result: dict[str, Any]) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, title=title)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 18)]
    story.append(Paragraph("LegalDoc AI Analysis Report", styles["Heading2"]))
    story.append(Spacer(1, 12))
    for heading, value in _flatten(result):
        story.append(Paragraph(heading, styles["Heading3"]))
        story.append(Paragraph(value.replace("\n", "<br/>"), styles["BodyText"]))
        story.append(Spacer(1, 10))
    doc.build(story)
    return buffer.getvalue()


def analysis_docx(*, title: str, result: dict[str, Any]) -> bytes:
    document = DocxDocument()
    document.add_heading(title, level=0)
    document.add_paragraph("LegalDoc AI Analysis Report")
    for heading, value in _flatten(result):
        document.add_heading(heading, level=1)
        document.add_paragraph(value)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
